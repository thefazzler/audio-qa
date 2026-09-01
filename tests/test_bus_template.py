"""The BUS Writing Template extractor, on a generated document and a real one.

Two layers, for the reason D16 gives. The structural tests build a document
from nothing, so they run on a fresh clone and prove the rules: SCRIPT only,
scene headers stripped but kept, placeholders dropped, blocks matched to
delivered files by order, a wrong COURSE ID halted. The golden tests run
against a real CGT delivery and are skipped when it is absent, because a real
script is narration text and narration text is customer material.

The generator is the committed part. Nothing in this file quotes a real script.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from qa.extract_docx import (
    build_freeform_topic,
    build_script_docx_bus,
    placeholder_reason,
    read_blocks,
)
from qa.util import ScriptError

# The real delivery, filed the way Courses 10 and 11 are: inside tests/, and
# gitignored. See HANDOVER.md.
REAL = Path(__file__).parent / "spcrisc26" / "course02"


def real_document() -> Path | None:
    if not REAL.is_dir():
        return None
    found = sorted(REAL.glob("*.docx"))
    return found[0] if found else None


# ---------------------------------------------------------------------------
# A BUS document from nothing
# ---------------------------------------------------------------------------

def make_bus_document(
    path: Path,
    course_id: str = "it_gen01_02",
    blocks: list[dict] | None = None,
    pronunciation: list[tuple[str, str]] = (),
) -> Path:
    """Write a document with the template's structure and invented narration.

    Mirrors what a real BUS file does: a course header table, a pronunciation
    guide, then per topic a single-cell heading, a single-cell metadata row and
    a SPEAKER/SCRIPT/OST table whose SCRIPT cells open with a bold scene header.
    """
    from docx import Document

    blocks = blocks if blocks is not None else [
        {
            "title": "COURSE OVERVIEW",
            "words": 12,
            "scenes": [("Scene 1: Learning Objectives", "The kettle boils a fresh pot.")],
        },
        {
            "title": "TOPIC 1 TITLE: REPOTTING THE CUTTINGS",
            "words": 20,
            "scenes": [
                ("Scene 1: Why Repot", "Volunteers repot the tallest cuttings."),
                ("Scene 2: When To Repot", "The greenhouse warms by noon each day."),
            ],
        },
        {
            "title": "TOPIC 2 TITLE: COURSE SUMMARY",
            "words": 9,
            "scenes": [("Scene 1: Summary", "The pot was repotted and the kettle boiled.")],
        },
    ]

    document = Document()

    header = document.add_table(rows=0, cols=2)
    for key, value in (
        ("CONTENT VERSION", "v1.0"),
        ("COURSE TITLE", "A generated course"),
        ("COURSE ID", course_id),
        ("TOTAL SCRIPT WORD COUNT", "41"),
        ("ESTIMATED COURSE DURATION (140 WPM)", "0m 18s"),
    ):
        row = header.add_row()
        row.cells[0].text = key
        row.cells[1].text = value

    guide = document.add_table(rows=2, cols=4)
    for cell in guide.rows[0].cells:
        cell.text = "Pronunciation Guide (add rows as needed)"
    for cell, label in zip(guide.rows[1].cells, ("Term", "Pronunciation", "Source (Link)", "Topic")):
        cell.text = label
    for term, say in pronunciation:
        row = guide.add_row()
        row.cells[0].text = term
        row.cells[1].text = say
        row.cells[2].text = ""
        row.cells[3].text = "1"

    for block in blocks:
        heading = document.add_table(rows=1, cols=1)
        heading.rows[0].cells[0].text = block["title"]

        metadata = document.add_table(rows=1, cols=1)
        metadata.rows[0].cells[0].text = (
            "VIDEO FILENAME: -\n"
            f"TOPIC WORD COUNT / TIME: {block['words']} / 0m 06s\n"
            "(USE METRIC 140 WORDS = 1 MINUTE)\n"
            f"OBJECTIVE: {block.get('objective', 'Do the thing')}"
        )

        script = document.add_table(rows=1, cols=3)
        for cell, label in zip(script.rows[0].cells, ("SPEAKER", "SCRIPT", "OST")):
            cell.text = label
        for scene, narration in block["scenes"]:
            row = script.add_row()
            row.cells[0].text = "HOST"
            cell = row.cells[1]
            cell.paragraphs[0].add_run(scene).bold = True
            cell.add_paragraph(" " + narration)
            row.cells[2].text = "A bullet nobody reads aloud\nAnother such bullet"

    document.add_paragraph("AI Generated Course Outline", style="Heading 2")
    document.add_paragraph("1. Course Overview")

    document.save(str(path))
    return path


@pytest.fixture
def generated(tmp_path) -> Path:
    return make_bus_document(tmp_path / "generated_scripts.docx")


# ---------------------------------------------------------------------------
# Structure
# ---------------------------------------------------------------------------

def test_every_block_is_found_in_document_order(generated):
    from qa.extract_docx import _open

    blocks, header, guide = read_blocks(_open(generated))
    assert [b.title for b in blocks] == [
        "COURSE OVERVIEW",
        "TOPIC 1 TITLE: REPOTTING THE CUTTINGS",
        "TOPIC 2 TITLE: COURSE SUMMARY",
    ]
    assert header["course id"] == "it_gen01_02"
    assert guide == []


def test_narration_is_the_script_column_and_never_the_ost(generated):
    """OST is on-screen bullet text. Reading it would invent a wall of deletions."""
    script = build_script_docx_bus(generated, ["01", "02", "03"], {}, "it_gen01_02_enus")
    narration = " ".join(s for t in script["topics"] for s in t["sentences"])
    assert "bullet nobody reads aloud" not in narration
    assert "Volunteers repot the tallest cuttings." in narration


def test_scene_headers_are_stripped_but_kept_with_their_position(generated):
    """A header the voice does read must surface as an insertion, not vanish."""
    script = build_script_docx_bus(generated, ["01", "02", "03"], {}, "it_gen01_02_enus")
    topic = script["topics"][1]
    assert [span["text"] for span in topic["non_narration"]] == [
        "Scene 1: Why Repot",
        "Scene 2: When To Repot",
    ]
    assert [span["sentence_index"] for span in topic["non_narration"]] == [0, 1]
    assert all(span["kind"] == "scene_header" for span in topic["non_narration"])
    assert not any("Scene 1" in s for s in topic["sentences"])


def test_multiple_speaker_rows_are_still_one_topic(generated):
    script = build_script_docx_bus(generated, ["01", "02", "03"], {}, "it_gen01_02_enus")
    assert script["topics"][1]["sentences"] == [
        "Volunteers repot the tallest cuttings.",
        "The greenhouse warms by noon each day.",
    ]


def test_the_author_word_count_and_estimate_are_carried(generated):
    script = build_script_docx_bus(generated, ["01", "02", "03"], {}, "it_gen01_02_enus")
    assert script["topics"][1]["author_word_count"] == 20
    assert script["topics"][1]["author_estimate"] == "0m 06s"
    assert script["topics"][1]["objective"] == "Do the thing"


def test_the_trailing_outline_section_is_not_a_topic(generated):
    script = build_script_docx_bus(generated, ["01", "02", "03"], {}, "it_gen01_02_enus")
    narration = " ".join(s for t in script["topics"] for s in t["sentences"])
    assert "AI Generated Course Outline" not in narration


def test_the_pronunciation_guide_becomes_watchlist_candidates(tmp_path):
    path = make_bus_document(
        tmp_path / "guided.docx", pronunciation=[("CRISC", "see-risk")]
    )
    script = build_script_docx_bus(path, ["01", "02", "03"], {}, "it_gen01_02_enus")
    assert script["pronunciation_guide"] == [
        {"term": "CRISC", "say": "see-risk", "source": "", "topic": "1"}
    ]


# ---------------------------------------------------------------------------
# Placeholders
# ---------------------------------------------------------------------------

PLACEHOLDER = {
    "title": "TOPIC 2 TITLE: A TAB ACTIVITY (HTML INTERACTIVITY)",
    "words": 8,
    "scenes": [("Scene 1: A Tab Activity", "This will be replaced by an HTML Interactivity")],
}


def test_a_placeholder_block_is_dropped_and_recorded(tmp_path):
    """Nothing is delivered for it, so it must not consume a delivered file."""
    blocks = [
        {"title": "COURSE OVERVIEW", "words": 6, "scenes": [("Scene 1: Intro", "The kettle boils.")]},
        PLACEHOLDER,
        {"title": "TOPIC 3 TITLE: COURSE SUMMARY", "words": 4, "scenes": [("Scene 1: Summary", "The pot was repotted.")]},
    ]
    path = make_bus_document(tmp_path / "placeholder.docx", blocks=blocks)
    script = build_script_docx_bus(path, ["01", "02"], {}, "it_gen01_02_enus")

    dropped = script["mapping"]["dropped_blocks"]
    assert len(dropped) == 1
    assert dropped[0]["block"] == 2
    assert "will be replaced by" in dropped[0]["reason"]
    assert len(script["topics"]) == 2


def test_a_short_interactivity_block_is_dropped_on_its_title(tmp_path):
    """The second signal, for a placeholder whose sentence was reworded."""
    from qa.extract_docx import _open

    blocks = [
        {"title": "COURSE OVERVIEW", "words": 6, "scenes": [("Scene 1: Intro", "The kettle boils.")]},
        {
            "title": "TOPIC 2 TITLE: A TAB ACTIVITY (HTML INTERACTIVITY)",
            "words": 5,
            "scenes": [("Scene 1: A Tab Activity", "Placeholder for the activity.")],
        },
    ]
    path = make_bus_document(tmp_path / "short.docx", blocks=blocks)
    found, _, _ = read_blocks(_open(path))
    assert placeholder_reason(found[1])
    assert not placeholder_reason(found[0])


def test_a_short_block_without_the_title_hint_is_kept(tmp_path):
    """Being brief is not being a placeholder. A summary is short and real."""
    from qa.extract_docx import _open

    path = make_bus_document(tmp_path / "brief.docx")
    found, _, _ = read_blocks(_open(path))
    assert all(not placeholder_reason(block) for block in found)


# ---------------------------------------------------------------------------
# Mapping, and the two halts
# ---------------------------------------------------------------------------

def test_blocks_map_to_delivered_files_by_order_not_by_heading_number(tmp_path):
    """Heading numbers include the placeholder; the delivered files do not."""
    blocks = [
        {"title": "COURSE OVERVIEW", "words": 6, "scenes": [("Scene 1: Intro", "The kettle boils.")]},
        {"title": "TOPIC 1 TITLE: FIRST", "words": 6, "scenes": [("Scene 1: One", "Volunteers repot the cuttings.")]},
        PLACEHOLDER,
        {"title": "TOPIC 3 TITLE: COURSE SUMMARY", "words": 4, "scenes": [("Scene 1: Summary", "The pot was repotted.")]},
    ]
    path = make_bus_document(tmp_path / "ordered.docx", blocks=blocks)
    script = build_script_docx_bus(path, ["01", "02", "03"], {}, "it_gen01_02_enus")

    pairs = {p["topic"]: p["block"] for p in script["mapping"]["pairs"]}
    assert pairs == {"01": 1, "02": 2, "03": 4}
    assert script["topics"][2]["heading_number"] == 3


def test_a_block_count_that_disagrees_with_the_delivery_halts_with_evidence(generated):
    with pytest.raises(ScriptError) as exc:
        build_script_docx_bus(generated, ["01", "02"], {}, "it_gen01_02_enus")
    message = str(exc.value)
    assert "PROBABLE MAPPING ERROR" in message
    assert "3 narrated blocks" in message
    assert "COURSE OVERVIEW" in message


def test_a_course_id_that_names_another_course_halts(generated):
    with pytest.raises(ScriptError, match="different course"):
        build_script_docx_bus(generated, ["01", "02", "03"], {}, "it_other99_07_enus")


def test_the_course_id_matches_the_filename_code_without_its_locale(generated):
    """COURSE ID carries no locale segment; the delivered filenames do."""
    script = build_script_docx_bus(generated, ["01", "02", "03"], {}, "it_gen01_02_enus")
    assert script["course_id"] == "it_gen01_02"


# ---------------------------------------------------------------------------
# Freeform
# ---------------------------------------------------------------------------

def test_a_freeform_text_file_is_all_narration(tmp_path):
    path = tmp_path / "demo_script.txt"
    path.write_text(
        "  The kettle boils a fresh pot. Volunteers repot the cuttings.  ",
        encoding="utf-8",
    )
    entry = build_freeform_topic("09", path)
    assert entry["script"] == "freeform"
    assert entry["scripted"] is True
    assert entry["sentences"] == [
        "The kettle boils a fresh pot.",
        "Volunteers repot the cuttings.",
    ]
    assert entry["source_ref"] == "demo_script.txt"


def test_a_freeform_docx_is_all_narration(tmp_path):
    from docx import Document

    document = Document()
    document.add_paragraph("The kettle boils a fresh pot.")
    document.add_paragraph("Volunteers repot the cuttings.")
    path = tmp_path / "demo_script.docx"
    document.save(str(path))

    entry = build_freeform_topic("09", path)
    assert entry["sentences"] == [
        "The kettle boils a fresh pot.",
        "Volunteers repot the cuttings.",
    ]


def test_an_empty_freeform_document_is_refused(tmp_path):
    path = tmp_path / "empty.txt"
    path.write_text("   \n", encoding="utf-8")
    with pytest.raises(ScriptError, match="is empty"):
        build_freeform_topic("09", path)


def test_a_freeform_script_in_an_unsupported_format_is_refused(tmp_path):
    path = tmp_path / "script.pdf"
    path.write_bytes(b"%PDF")
    with pytest.raises(ScriptError, match="must be .docx or .txt"):
        build_freeform_topic("09", path)


# ---------------------------------------------------------------------------
# Golden assertions against the real delivery
# ---------------------------------------------------------------------------

realonly = pytest.mark.skipif(
    real_document() is None,
    reason=(
        "no CGT script document in tests/spcrisc26/course02; it is customer "
        "material and is not in version control"
    ),
)

# The delivered order the fixture implies: the overview is topic 1, then the
# document's TOPIC 1 to TOPIC 10 with the placeholder removed.
REAL_TOPICS = [f"{n:02d}" for n in range(1, 11)]


@pytest.fixture
def real_script():
    return build_script_docx_bus(
        real_document(), REAL_TOPICS, {}, "it_spcrisc26_02_enus"
    )


@realonly
def test_real_document_has_eleven_blocks_and_drops_one(real_script):
    assert real_script["mapping"]["blocks_found"] == 11
    assert len(real_script["mapping"]["dropped_blocks"]) == 1
    assert len(real_script["topics"]) == 10


@realonly
def test_the_dropped_block_is_the_interactivity_placeholder(real_script):
    dropped = real_script["mapping"]["dropped_blocks"][0]
    assert dropped["block"] == 10
    assert "INTERACTIVITY" in dropped["title"].upper()
    assert dropped["word_count"] < 20


@realonly
def test_the_longest_topic_is_the_demo_and_carries_no_ost(real_script):
    """The document's TOPIC 6 is the demo: about 1,600 words, SCRIPT only.

    The strongest evidence that no OST leaked in is that every block's word
    count equals the count the author wrote in its metadata row. OST would add
    hundreds of words per topic.
    """
    longest = max(real_script["topics"], key=lambda t: t["word_count"])
    assert 1500 < longest["word_count"] < 1700
    assert longest["word_count"] == longest["author_word_count"]


@realonly
def test_every_author_word_count_matches_what_was_extracted(real_script):
    mismatched = [
        (t["topic"], t["word_count"], t["author_word_count"])
        for t in real_script["topics"]
        if t["word_count"] != t["author_word_count"]
    ]
    assert not mismatched


@realonly
def test_every_scene_header_is_captured_as_non_narration(real_script):
    spans = [s for t in real_script["topics"] for s in t["non_narration"]]
    assert len(spans) >= 30
    assert all(s["kind"] == "scene_header" for s in spans)
    assert all(s["bold"] for s in spans)
    narration = " ".join(s for t in real_script["topics"] for s in t["sentences"])
    assert "Scene 1:" not in narration


@realonly
def test_the_course_id_is_the_courses_own(real_script):
    assert real_script["course_id"] == "it_spcrisc26_02"


@realonly
def test_the_pronunciation_guide_is_empty_in_this_delivery(real_script):
    """Empty here, and the extractor must say so rather than inventing rows."""
    assert real_script["pronunciation_guide"] == []


# ---------------------------------------------------------------------------
# What reaches the packet
# ---------------------------------------------------------------------------

def test_the_authors_word_count_and_estimate_reach_the_packet(generated):
    """A pacing reference from the source side, and never a threshold."""
    from qa.packet import _author_estimates

    script = build_script_docx_bus(generated, ["01", "02", "03"], {}, "it_gen01_02_enus")
    lines = _author_estimates(script)
    text = "\n".join(lines)

    assert "Author's word count" in text
    assert "0m 06s" in text
    assert "reference" in text and "threshold" in text
    # One row per topic, plus the header rows.
    assert sum(1 for line in lines if line.startswith("| 0")) == 3


def test_a_storyboard_course_gets_no_author_estimate_section(tmp_path):
    """A pptx carries no such number, so the section is simply absent."""
    from qa.packet import _author_estimates

    assert _author_estimates({"topics": [{"topic": "01", "word_count": 10}]}) == []
    assert _author_estimates({}) == []


def test_a_disagreement_with_the_authors_count_is_shown_not_hidden(generated):
    """If extraction and the author disagree, the reader sees by how much."""
    from qa.packet import _author_estimates

    script = build_script_docx_bus(generated, ["01", "02", "03"], {}, "it_gen01_02_enus")
    script["topics"][0]["word_count"] = script["topics"][0]["author_word_count"] + 4
    text = "\n".join(_author_estimates(script))
    assert "(+4)" in text


def test_the_dropped_block_reaches_the_packet_with_its_reason(tmp_path):
    from qa.packet import _dropped_blocks

    blocks = [
        {"title": "COURSE OVERVIEW", "words": 6, "scenes": [("Scene 1: Intro", "The kettle boils.")]},
        PLACEHOLDER,
        {"title": "TOPIC 3 TITLE: COURSE SUMMARY", "words": 4, "scenes": [("Scene 1: Summary", "The pot was repotted.")]},
    ]
    path = make_bus_document(tmp_path / "placeholder.docx", blocks=blocks)
    script = build_script_docx_bus(path, ["01", "02"], {}, "it_gen01_02_enus")

    text = "\n".join(_dropped_blocks(script))
    assert "not treated as a topic" in text
    assert "HTML INTERACTIVITY" in text.upper()
    assert "will be replaced by" in text
