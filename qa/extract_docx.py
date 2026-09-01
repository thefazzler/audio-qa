"""Script extraction from Word documents: the BUS template, and freeform.

Two extractors, both emitting the per-topic structure the aligner already
consumes, so nothing downstream of the script stage knows which one ran.

`docx_bus` reads the "BUS Writing Template", which is how a CGT course carries
its narration. There is no PowerPoint on a CGT course at all, so the storyboard
extractor has nothing to work with and this is not an alternative path but the
only one.

`freeform` reads a whole document as narration, for the occasional vendor demo
that arrives with a script of its own.

What this module refuses to do is infer. The template's structure was read off
a real delivery and every rule below is checked against it; where the document
does not say something, this raises rather than guessing. A script extractor
that guesses wrong produces a page of false deletions that reads exactly like a
narration catastrophe, which is the most expensive mistake in the pipeline.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from .util import ScriptError, split_sentences

# --------------------------------------------------------------------------
# Recognizing the parts of the template
# --------------------------------------------------------------------------

# A block opens with a single-cell header row. The course overview has no topic
# number; every other block does, including the summary, which arrives as
# "TOPIC 10 TITLE: COURSE SUMMARY" rather than under a heading of its own.
BLOCK_HEADER = re.compile(
    r"^\s*(?:COURSE\s+OVERVIEW|COURSE\s+SUMMARY|TOPIC\s+(\d+)\s+TITLE\s*:)",
    re.IGNORECASE,
)

# The metadata row that follows it. Both wordings seen in the fixture:
# "TOPIC WORD COUNT / TIME: 863 / 6m 10s" and, on the overview,
# "COURSE OVERVIEW WORD COUNT (60-100 WORDS): 138 / 0m 59s".
METADATA_ROW = re.compile(r"WORD\s+COUNT", re.IGNORECASE)
WORD_COUNT = re.compile(
    r"WORD\s+COUNT[^:\n]*:\s*(\d+)\s*(?:/\s*([^\n]+))?", re.IGNORECASE
)
OBJECTIVE = re.compile(r"^\s*OBJECTIVE\s*:\s*(.+)$", re.IGNORECASE | re.MULTILINE)

# The script table's own header row. This is the anchor: a table with these
# three columns is a script table and nothing else in the document is.
SCRIPT_COLUMNS = ("speaker", "script", "ost")

# A scene header inside a SCRIPT cell. Written bold at the start of its own
# paragraph, and not narrated.
SCENE_HEADER = re.compile(r"^\s*Scene\s+\d+\s*:\s*\S", re.IGNORECASE)

# The pronunciation guide, which is watchlist input when it is filled in.
PRONUNCIATION_HEADER = re.compile(r"Pronunciation\s+Guide", re.IGNORECASE)
PRONUNCIATION_COLUMNS = ("term", "pronunciation", "source", "topic")

# The course header table, a two column list of properties.
COURSE_ID_KEY = "course id"

# --------------------------------------------------------------------------
# Placeholder blocks
# --------------------------------------------------------------------------
# A block that stands in for an interactivity is not narrated and nothing is
# delivered for it, so it must not consume one of the delivered files. Two
# independent signals, because either alone would be brittle: the template's
# own placeholder sentence, and a block too short to be a topic whose title
# says what it is. Both fire on the fixture's topic 9; either is enough.

PLACEHOLDER_SENTENCE = ("will be replaced by", "interactivity")
PLACEHOLDER_TITLE_HINT = "INTERACTIVITY"
PLACEHOLDER_WORD_LIMIT = 20

# The trailing section the authoring tool appends after the last topic. It is
# not script and it is not narrated.
TRAILING_SECTION = "AI Generated Course Outline"


@dataclass
class Block:
    """One topic-shaped block of the template, before placeholders are dropped."""

    index: int
    title: str
    heading_number: int | None
    author_word_count: int | None = None
    author_estimate: str = ""
    objective: str = ""
    sentences: list[str] = field(default_factory=list)
    sentence_rows: list[int] = field(default_factory=list)
    non_narration: list[dict] = field(default_factory=list)
    speakers: list[str] = field(default_factory=list)
    raw_script: str = ""

    @property
    def word_count(self) -> int:
        return sum(len(s.split()) for s in self.sentences)


def _open(path: Path):
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise ScriptError(
            "Reading a Word script needs python-docx, which is not installed.\n"
            "  Install it with:  pip install -e .\n"
            f"  ({exc})"
        ) from exc
    try:
        return Document(str(path))
    except Exception as exc:  # python-docx raises a variety of package errors
        raise ScriptError(f"Could not open {path.name}:\n  {exc}") from exc


def _cells(row) -> list[str]:
    return [cell.text.strip() for cell in row.cells]


def _is_single_cell(table) -> bool:
    """One row whose cells are all the same merged cell."""
    if len(table.rows) != 1:
        return False
    texts = {cell.text for cell in table.rows[0].cells}
    return len(texts) == 1


def _is_script_table(table) -> bool:
    if not table.rows:
        return False
    header = [text.lower() for text in _cells(table.rows[0])]
    return len(header) >= 3 and tuple(header[:3]) == SCRIPT_COLUMNS


def _is_pronunciation_table(table) -> bool:
    for row in table.rows[:2]:
        cells = [text.lower() for text in _cells(row)]
        if len(cells) >= 4 and tuple(c.split(" ")[0] for c in cells[:4]) == (
            PRONUNCIATION_COLUMNS
        ):
            return True
    return False


# --------------------------------------------------------------------------
# The SCRIPT cell
# --------------------------------------------------------------------------

def _bold_prefix(paragraph) -> str:
    """The run text at the start of a paragraph that is set bold."""
    text = ""
    for run in paragraph.runs:
        if not run.bold:
            break
        text += run.text
    return text


def read_script_cell(cell, row_index: int) -> tuple[list[str], list[dict], str]:
    """Narration sentences and the scene headers stripped out of them.

    Scene headers are not narrated, so leaving them in the alignment text would
    report every one of them as a deletion. They are kept as tagged
    non-narration spans instead, with the position they held, so that a header
    the voice *did* read surfaces as an insertion rather than disappearing.
    """
    sentences: list[str] = []
    headers: list[dict] = []
    raw: list[str] = []

    for paragraph in cell.paragraphs:
        text = paragraph.text
        raw.append(text)
        stripped = text.strip()
        if not stripped:
            continue

        bold = _bold_prefix(paragraph).strip()
        header = ""
        remainder = stripped
        if bold and SCENE_HEADER.match(bold):
            header = bold
            remainder = stripped[len(bold):].strip() if stripped.startswith(bold) else ""
        elif SCENE_HEADER.match(stripped):
            # The pattern without the bold. Treated as a header anyway, and
            # reported, because a header left in the narration would be read as
            # a deletion at every scene in the course.
            header = stripped
            remainder = ""

        if header:
            headers.append(
                {
                    "kind": "scene_header",
                    "text": header,
                    "row": row_index,
                    "sentence_index": len(sentences),
                    "bold": bool(bold and SCENE_HEADER.match(bold)),
                }
            )
        if remainder:
            sentences.extend(split_sentences(remainder))

    return sentences, headers, "\n".join(raw)


# --------------------------------------------------------------------------
# Walking the document
# --------------------------------------------------------------------------

def read_blocks(document) -> tuple[list[Block], dict, list[dict]]:
    """Every topic block in document order, plus the course header and guide.

    One pass over the tables. A block opens on a single-cell header row, takes
    the metadata row that follows, and closes on its SPEAKER/SCRIPT/OST table.
    Everything else in the document, including the trailing outline section, is
    simply not one of those three things and is ignored.
    """
    blocks: list[Block] = []
    header_fields: dict = {}
    guide: list[dict] = []
    pending: Block | None = None

    for table in document.tables:
        if _is_pronunciation_table(table):
            guide.extend(_read_pronunciation(table))
            continue

        if not header_fields:
            header_fields = _read_course_header(table)
            if header_fields:
                continue

        if _is_script_table(table):
            if pending is None:
                # A script table with no block header above it. The document is
                # not shaped the way this extractor was written against, and
                # attaching its rows to the previous topic would silently move
                # narration between topics.
                raise ScriptError(
                    "A SPEAKER/SCRIPT/OST table appears before any topic "
                    "heading in this document. Every script table must sit "
                    "under a 'COURSE OVERVIEW' or 'TOPIC N TITLE:' row."
                )
            _read_script_table(table, pending)
            blocks.append(pending)
            pending = None
            continue

        if not _is_single_cell(table):
            continue

        text = table.rows[0].cells[0].text
        match = BLOCK_HEADER.match(text)
        if match:
            if pending is not None:
                raise ScriptError(
                    f"Topic heading {_one_line(text)!r} follows "
                    f"{pending.title!r} with no script table between them."
                )
            pending = Block(
                index=len(blocks),
                title=_one_line(text),
                heading_number=int(match.group(1)) if match.group(1) else None,
            )
            continue

        if pending is not None and METADATA_ROW.search(text):
            _read_metadata(text, pending)

    if pending is not None:
        raise ScriptError(
            f"Topic heading {pending.title!r} has no SPEAKER/SCRIPT/OST table "
            "after it, so the document says nothing about what is narrated there."
        )
    return blocks, header_fields, guide


def _one_line(text: str) -> str:
    return " ".join(text.split())


def _read_course_header(table) -> dict:
    """The two column property list at the top, if this table is it."""
    fields: dict = {}
    for row in table.rows:
        cells = _cells(row)
        if len(cells) < 2:
            return {}
        fields[cells[0].strip().lower()] = cells[1].strip()
    return fields if COURSE_ID_KEY in fields else {}


def _read_pronunciation(table) -> list[dict]:
    """Rows of the pronunciation guide, skipping its two header rows."""
    entries: list[dict] = []
    for row in table.rows:
        cells = _cells(row)
        if len(cells) < 4:
            continue
        term = cells[0].strip()
        if not term or term.lower() == "term" or PRONUNCIATION_HEADER.search(term):
            continue
        entries.append(
            {
                "term": term,
                "say": cells[1].strip(),
                "source": cells[2].strip(),
                "topic": cells[3].strip(),
            }
        )
    return entries


def _read_metadata(text: str, block: Block) -> None:
    counts = WORD_COUNT.search(text)
    if counts:
        block.author_word_count = int(counts.group(1))
        block.author_estimate = (counts.group(2) or "").strip()
    objective = OBJECTIVE.search(text)
    if objective:
        block.objective = objective.group(1).strip()


def _read_script_table(table, block: Block) -> None:
    """Every SCRIPT cell in the table, in order, as this block's narration.

    The SCRIPT column only. OST is on-screen bullet text that nobody reads
    aloud, and including it would produce a wall of false deletions in every
    topic of every CGT course.
    """
    raw: list[str] = []
    for index, row in enumerate(table.rows):
        cells = _cells(row)
        if index == 0 and len(cells) >= 3 and cells[1].lower() == "script":
            continue
        if len(row.cells) < 2:
            continue
        sentences, headers, text = read_script_cell(row.cells[1], index)
        for header in headers:
            header["sentence_index"] += len(block.sentences)
            block.non_narration.append(header)
        block.sentences.extend(sentences)
        block.sentence_rows.extend([index] * len(sentences))
        block.speakers.append(cells[0].strip())
        raw.append(text)
    block.raw_script = "\n".join(raw)


# --------------------------------------------------------------------------
# Placeholders
# --------------------------------------------------------------------------

def placeholder_reason(block: Block) -> str:
    """Why this block is a stand-in rather than a topic, or "" if it is one."""
    lowered = " ".join(block.sentences).lower()
    if all(hint in lowered for hint in PLACEHOLDER_SENTENCE):
        return (
            "the block's script is the template's interactivity placeholder "
            f'sentence ("{PLACEHOLDER_SENTENCE[0]} ... {PLACEHOLDER_SENTENCE[1]}")'
        )
    if (
        block.word_count < PLACEHOLDER_WORD_LIMIT
        and PLACEHOLDER_TITLE_HINT in block.title.upper()
    ):
        return (
            f"the block has {block.word_count} words, below "
            f"{PLACEHOLDER_WORD_LIMIT}, and its title says "
            f"{PLACEHOLDER_TITLE_HINT.lower()}"
        )
    return ""


# --------------------------------------------------------------------------
# The extractor
# --------------------------------------------------------------------------

def build_script_docx_bus(
    document_path: Path,
    topics: list[str],
    topic_scripts: dict | None = None,
    course_code: str = "",
) -> dict:
    """Extract per-topic narration from a BUS Writing Template document."""
    document = _open(document_path)
    blocks, header_fields, guide = read_blocks(document)

    dropped: list[dict] = []
    kept: list[Block] = []
    for block in blocks:
        reason = placeholder_reason(block)
        if reason:
            dropped.append(
                {
                    "block": block.index + 1,
                    "title": block.title,
                    "reason": reason,
                    "word_count": block.word_count,
                    "author_word_count": block.author_word_count,
                }
            )
        else:
            kept.append(block)

    course_id = (header_fields.get(COURSE_ID_KEY) or "").strip()
    warnings = _check_course_id(course_id, course_code, document_path)

    if len(kept) != len(topics):
        raise ScriptError(
            _mapping_error(document_path, blocks, kept, dropped, topics)
        )

    unbolded = [
        span
        for block in kept
        for span in block.non_narration
        if not span.get("bold")
    ]
    if unbolded:
        warnings.append(
            f"{len(unbolded)} scene header"
            f"{'' if len(unbolded) == 1 else 's'} matched the 'Scene N:' "
            "pattern without being bold. They were treated as headers and kept "
            "out of the alignment text; if any of them is really narration, it "
            "will now read as a deletion."
        )

    topic_entries: list[dict] = []
    scripts = topic_scripts or {}
    for topic, block in zip(topics, kept):
        state = scripts.get(topic)
        entry = {
            "topic": topic,
            "script": getattr(state, "state", "verbatim"),
            "scripted": getattr(state, "aligned", True),
            "slides": None,
            "source_ref": block.title,
            "block": block.index + 1,
            "heading_number": block.heading_number,
            "objective": block.objective,
            "author_word_count": block.author_word_count,
            "author_estimate": block.author_estimate,
            "speakers": sorted({s for s in block.speakers if s}),
            "sentences": block.sentences,
            "sentence_rows": block.sentence_rows,
            "non_narration": block.non_narration,
            "word_count": block.word_count,
        }
        topic_entries.append(entry)

    return {
        "script_source": "docx_bus",
        "script_document": document_path.name,
        "storyboard": None,
        "slide_count": None,
        "course_id": course_id,
        "course_title": header_fields.get("course title", ""),
        "author_total_word_count": _as_int(
            header_fields.get("total script word count")
        ),
        "author_total_estimate": header_fields.get(
            "estimated course duration (140 wpm)", ""
        ),
        "pronunciation_guide": guide,
        "mapping": {
            "source": "docx_bus, blocks in document order",
            "blocks_found": len(blocks),
            "blocks_used": len(kept),
            "dropped_blocks": dropped,
            "markers": [],
            "excluded_slides": [],
            "pairs": [
                {"topic": topic, "block": block.index + 1, "title": block.title}
                for topic, block in zip(topics, kept)
            ],
        },
        "warnings": warnings,
        "topics": topic_entries,
    }


def _as_int(value: object) -> int | None:
    try:
        return int(str(value).strip())
    except (TypeError, ValueError):
        return None


def _check_course_id(course_id: str, course_code: str, path: Path) -> list[str]:
    """The document must be this course's document.

    The COURSE ID cell carries the code without the locale segment, so the
    comparison drops the locale from the filename-derived code rather than
    demanding they be identical.
    """
    if not course_code:
        return []
    if not course_id:
        return [
            f"{path.name} has no COURSE ID in its header table, so the document "
            "could not be checked against the delivered filenames."
        ]
    expected = course_code.rstrip("_").lower()
    without_locale = "_".join(expected.split("_")[:-1])
    if course_id.lower() in {expected, without_locale}:
        return []
    raise ScriptError(
        f"The script document is for a different course.\n"
        f"  {path.name} says COURSE ID {course_id!r}\n"
        f"  the delivered filenames say {course_code!r}\n"
        "  Aligning a course against another course's script would report every "
        "topic as a total narration failure. Check which document belongs here."
    )


def _mapping_error(
    path: Path,
    blocks: list[Block],
    kept: list[Block],
    dropped: list[dict],
    topics: list[str],
) -> str:
    lines = [
        "PROBABLE MAPPING ERROR: "
        f"{path.name} carries {len(kept)} narrated blocks but the course "
        f"delivered {len(topics)} audio files.",
        "",
        "  Blocks found, in document order:",
    ]
    dropped_indices = {d["block"] for d in dropped}
    for block in blocks:
        mark = "dropped " if block.index + 1 in dropped_indices else "        "
        lines.append(
            f"    {block.index + 1:>3}  {mark}{block.word_count:>5} words  {block.title[:70]}"
        )
    if dropped:
        lines += ["", "  Dropped as placeholders:"]
        for item in dropped:
            lines.append(f"    {item['block']:>3}  {item['reason']}")
    lines += [
        "",
        f"  Audio topics: {', '.join(topics)}",
        "",
        "  Fix one of these:",
        "    - check that every narration file was delivered",
        "    - check whether a block that is really a placeholder is not being "
        "recognized as one, or the reverse",
        "    - set the per-topic script state in course.yaml if a delivered "
        "topic has no block in this document",
    ]
    return "\n".join(lines)


# --------------------------------------------------------------------------
# Freeform
# --------------------------------------------------------------------------

FREEFORM_SUFFIXES = {".docx", ".txt"}


def read_freeform(path: Path) -> str:
    """The whole document as narration. No structure is inferred at all."""
    suffix = path.suffix.lower()
    if suffix not in FREEFORM_SUFFIXES:
        allowed = " or ".join(sorted(FREEFORM_SUFFIXES))
        raise ScriptError(
            f"A freeform script must be {allowed}; {path.name} is {suffix or 'extensionless'}."
        )
    if suffix == ".txt":
        try:
            return path.read_text(encoding="utf-8").strip()
        except OSError as exc:
            raise ScriptError(f"Could not read {path.name}:\n  {exc}") from exc
        except UnicodeDecodeError as exc:
            raise ScriptError(
                f"{path.name} is not UTF-8 text:\n  {exc}\n"
                "  Save it as UTF-8, or supply it as a .docx."
            ) from exc

    document = _open(path)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts).strip()


def build_freeform_topic(topic: str, path: Path) -> dict:
    """One topic's entry, from a document that is entirely its narration."""
    text = read_freeform(path)
    if not text:
        raise ScriptError(
            f"{path.name} is empty, so there is no narration to align topic "
            f"{topic} against.\n"
            "  Set that topic's script to none if it really has no script."
        )
    sentences = split_sentences(text)
    return {
        "topic": topic,
        "script": "freeform",
        "scripted": True,
        "slides": None,
        "source_ref": path.name,
        "sentences": sentences,
        "sentence_rows": [0] * len(sentences),
        "non_narration": [],
        "word_count": sum(len(s.split()) for s in sentences),
    }
